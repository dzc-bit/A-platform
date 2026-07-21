# -*- coding: utf-8 -*-
"""生成 9 天实习日记（2026-07-17 ~ 2026-07-25），每天一个 docx，置于 D:/shixi/实习日记/。
格式严格对齐截图模板：标题 + 表格头（姓名/专业/学号/学院/单位/地点/日期）+ 正文两节 + 签名。"""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

CN = "宋体"
EN = "Times New Roman"

def set_font(run, size=12, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = EN
    rPr = run._element.get_or_add_rPr()
    ea = rPr.find(qn('w:eastAsia'))
    if ea is None:
        ea = OxmlElement('w:eastAsia'); rPr.append(ea)
    ea.set(qn('w:val'), CN)

def set_cell_font(cell, size=12):
    """设置表格内所有 run 的字体"""
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            set_font(r, size=size)

def para(doc, text="", size=12, bold=False, align=None, indent=False, before=0, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold)
    return p

def make_header_table(doc, date_str):
    """按截图模板创建信息表格：2列布局（姓名|专业 / 学号|学院 / 单位/地点/日期各占整行）。"""
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 设置表格整体宽度与边框样式
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    # 边框
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)

    def set_cell(row, col, text, colspan=1):
        cell = table.cell(row, col)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r, size=12)
        if colspan > 1:
            cell.merge(table.cell(row, col + 1))

    # Row 0: 姓　　名: xxx   |   专　　业: xxx
    set_cell(0, 0, "姓　　名：" + NAME)
    set_cell(0, 1, "专　　业：" + MAJOR)

    # Row 1: 学　　号: xxx   |   学　　院: xxx
    set_cell(1, 0, "学　　号：________")
    set_cell(1, 1, "学　　院：" + SCHOOL)

    # Row 2-4: 实习单位 / 实习地点 / 实习日期 —— 各跨两列占满整行
    set_cell(2, 0, "实习单位：" + UNIT, colspan=2)
    set_cell(3, 0, "实习地点：" + PLACE, colspan=2)
    set_cell(4, 0, "实习日期：                    " + date_str, colspan=2)

    # 统一设置所有单元格字体和内边距
    for row in table.rows:
        for cell in row.cells:
            set_cell_font(cell, size=12)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            # 设置单元格边距
            mar = OxmlElement('w:tcMar')
            for side in ['top', 'left', 'bottom', 'right']:
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), '80')   # 约 3pt 内边距
                m.set(qn('w:type'), 'dxa')
                mar.append(m)
            tcPr.append(mar)
    return table

# ============================================================
# 基本信息
# ============================================================
NAME = "邓子川"
MAJOR = "计算机科学与技术"
SCHOOL = "信息科学技术"
UNIT = "东软教育科技集团有限公司"   # 实训方案封面单位
PLACE = "暨南大学"                  # 集中实习地点

# ============================================================
# 9 天日记内容（已丰富，结合项目真实功能 & 实训方案主线）
# ============================================================
entries = [
    {
        "d": "7月17日", "date": "2026年7月17日",
        "work": (
            "上午参加实训启动会，导师详细介绍了本次实训的目标与考核标准，并发放了《暨南大学2023级实训方案3（基于Python的AI应用开发）》文档。"
            "明确了核心任务：以「东软智慧商务AI助手平台」为载体，完成从需求分析到部署上线的全流程开发。\n\n"
            "下午搭建开发环境：安装 Python 3.12、配置 VS Code 与 Cursor 编辑器插件、创建虚拟环境并安装 FastAPI、uvicorn、"
            "httpx 等基础依赖。克隆项目仓库后通读 README.md 和 docs/ 目录下的需求文档（requirements.md）、架构设计（architecture.md）"
            "以及 AI 设计方案（ai-design.md），初步理解了项目的分层架构——前端 Vue3 + 后端 FastAPI + AI 层 LangChain/LangGraph/Dify。\n\n"
            "尝试运行 `make dev` 一键启动命令，成功在本地跑通了前后端联调的 Hello World 接口，浏览器能看到 Vue3 页面正确调用 FastAPI 返回的数据。"
        ),
        "think": (
            "今天是实训第一天，最大的收获是对「AI 应用全链路」有了全局认知。以前只零散地用过一些 API，这次系统性地看到了从需求→架构→编码→部署的完整闭环。"
            "环境搭建看似琐碎，但 Python 虚拟环境、依赖版本锁定、前后端跨域配置这些细节如果不提前理顺，后面调试会非常痛苦。"
            "另外注意到项目 docs 写得很规范（需求/架构/AI设计/知识点对照各一份），这种工程化文档习惯值得在以后的项目中坚持。"
        ),
    },
    {
        "d": "7月18日", "date": "2026年7月18日",
        "work": (
            "全天聚焦于前后端基础功能的实现。\n\n"
            "后端方面：深入学习 FastAPI 框架的核心概念——路由定义与请求模型（Pydantic）、依赖注入（Depends）做权限校验、"
            "中间件记录请求日志、自定义异常处理器统一返回格式。使用 SQLAlchemy ORM 定义 User/BusinessData 数据表模型，"
            "实现了用户注册登录（JWT Token 签发与验证）和商务数据的 CRUD 接口（增删改查共 6 个端点）。\n\n"
            "前端方面：基于 Vue3 Composition API（setup 语法糖）开发了登录页和仪表盘首页。使用 ref/reactive 管理组件状态，"
            "props/emit 做父子组件通信，Vue Router 配置路由守卫（未登录跳转 login），Pinia 存储用户 token 与角色信息。"
            "页面 UI 使用了 Element Plus 组件库，表格带分页与搜索筛选功能。\n\n"
            "借助 Cursor/Codex AI 编码辅助工具生成了部分样板代码（如表单验证规则、API 请求封装），再人工审核调整业务逻辑。"
        ),
        "think": (
            "今天体会到前后端分离架构的优势——后端专注数据与接口，前端专注交互与状态，职责清晰、便于并行开发。"
            "Pinia 相比 Vuex3 确实更简洁，TypeScript 支持也更好。AI 辅助编码在生成 CRUD 样板代码时效率提升明显，但生成的代码往往缺少边界检查和异常处理，"
            "必须逐行审核才能合并。这再次说明 AI 是「副驾驶」，方向盘还在开发者手里。"
        ),
    },
    {
        "d": "7月19日", "date": "2026年7月19日",
        "work": (
            "进入大模型 API 集成阶段。\n\n"
            "首先学习 OpenAI 兼容 API 协议规范（/v1/chat/completions 接口的 request/response 字段含义）。使用 httpx 异步客户端分别对接了 "
            "DeepSeek-V3 和通义千问（Qwen-Max）两个大模型的 API，完成了对话补全接口的封装。重点实现了 SSE（Server-Sent Events）流式输出："
            "后端用 `StreamingResponse` 逐块转发 LLM 返回的 `data: {...}\n\n` 事件流，前端用 EventSource / Fetch Reader 实时接收并在聊天界面逐字渲染，"
            "体验接近 ChatGPT 的打字机效果。\n\n"
            "下午投入提示词工程（Prompt Engineering）实践。学习了提示词四要素框架——角色设定（Role）、任务描述（Task）、输出格式（Format）、约束条件（Constraint），"
            "并据此设计了三套提示词模板：（1）客服类——引导模型以礼貌专业的话术回答用户咨询；（2）知识查询类——要求模型基于检索到的上下文作答并标注引用来源；"
            "(3) 分析报告类——规定输出的 JSON 结构化字段。每套模板都经历了「编写→测试→评估质量→修改措辞→再测试」至少三轮迭代优化。"
        ),
        "think": (
            "流式输出对用户体验的提升是质的改变——等待整段生成完毕 vs 逐字呈现，心理感知延迟差距巨大。提示词工程的收获更大："
            "之前以为提示词就是「自然语言描述」，今天发现它更像一种「编程语言」——需要精确的语法（分隔符、JSON 模板）、严谨的逻辑（few-shot 示例引导）、"
            "以及系统的评估方法（同一组测试集对比修改前后的输出质量）。好的提示词能将模型可用率从 60% 提升到 90% 以上，这个杠杆效应令人印象深刻。"
        ),
    },
    {
        "d": "7月20日", "date": "2026年7月20日",
        "work": (
            "今天的学习重点是 LangChain 框架与 RAG（检索增强生成）技术。\n\n"
            "上午系统学习了 LangChain 的六大核心组件：Model I/O（模型输入输出封装）、Retrieval（检索模块含 DocumentLoader/TextSplitter/VectorStore/Retriever）、"
            "Chain（链式编排）、Memory（记忆机制含 BufferMemory/WindowMemory/SummaryMemory）、Tool（工具调用）以及 Agent（智能体决策）。"
            "动手实现了两条 Chain：一条用于客服问答（用户问题→意图识别→模板选择→LLM 生成），另一条用于知识查询（问题→检索上下文→注入 Prompt→生成带来源的回答）。\n\n"
            "下午搭建 RAG 检索流水线：使用 TextSplitter 将企业文档按 500 字符 + 50 重叠切分 chunks；调用 Embedding API 将每个 chunk 向量化并存入 FAISS 向量数据库；"
            "查询时先将用户问题向量化，再做相似度 Top-K 检索。在项目中进一步实现了「混合检索」策略：同时执行语义向量检索和 BM25 关键词检索，"
            "使用 RRF（Reciprocal Rank Fusion）加权融合两组结果，最终召回质量显著优于单一检索方式。回答中自动附带引用来源编号，用户可点击查看原文片段。"
        ),
        "think": (
            "RAG 解决了大模型「幻觉」和「知识过期」两大痛点——让模型基于企业私有数据作答而非凭空编造。今天的实践中发现几个关键经验："
            "(1) 分块大小直接影响检索效果——太大则噪声多，太小则语义不完整，500 字符左右对本项目文档比较合适；(2) "
            "混合检索（向量+关键词）比纯向量检索稳健得多，尤其是面对缩写、型号等精确匹配场景；(3) "
            "确定性 Embedding（对相同文本始终产出相同向量）是实现可复现检索的前提，项目中通过内容哈希缓存策略保证了这一点。"
        ),
    },
    {
        "d": "7月21日", "date": "2026年7月21日",
        "work": (
            "深入 LangGraph 多智能体编排与 Function Calling。\n\n"
            "理解了 Agent 的通用架构模式：感知（Perception，接收用户输入）→ 规划（Planning，拆解任务步骤）→ 工具调用（Tool Use，执行外部操作）"
            "→ 反馈（Observation，整合结果）→ 响应（Response，输出答案）。使用 LangGraph 的 StateGraph 定义了一个多 Agent 工作流，包含四个节点：\n\n"
            "• ClassificationAgent（分类节点）：判断用户意图属于「闲聊/知识查询/工单办理/数据分析」哪一类；\n"
            "• KnowledgeQueryAgent（知识检索节点）：调用 RAG 检索器获取相关文档片段；\n"
            "• ResponseAgent（回复生成节点）：根据分类结果和检索上下文组织最终回答；\n"
            "• BusinessAgentOrchestrator（业务编排节点）：处理工单创建、状态流转等结构化操作。\n\n"
            "节点之间通过条件边（Conditional Edge）路由——例如分类为「知识查询」时走 KnowledgeQueryAgent，分类为「工单」时走 BusinessAgent。"
            "同时掌握了 Function Calling 的完整流程：定义工具 Schema（JSON Schema 描述参数）→ 注册给 LLM → 解析 tool_calls → 执行函数 → 将结果回传 LLM 生成最终回答。"
            "特别注重安全措施：工具白名单限制、参数类型/枚举/范围校验、无来源证据时转人工而不臆造。"
        ),
        "think": (
            "多智能体的核心价值在于「可解释性」和「可控性」——复杂任务被拆分成独立节点，每个节点的输入输出都可观测、可调试、可单独替换。"
            "相比单一大模型 Prompt 搞定一切的方式，LangGraph 让 AI 应用的行为变得可预测、可维护。"
            "安全方面印象最深的是「工具白名单」原则——绝对不能让 LLM 自由执行任意代码或 SQL，必须限定在预先声明且参数经过严格校验的安全范围内。"
            "这是产业级 AI 应用与非玩具 Demo 的重要分水岭。"
        ),
    },
    {
        "d": "7月22日", "date": "2026年7月22日",
        "work": (
            "Dify 低代码平台的部署与应用集成。\n\n"
            "上午在 Docker 环境中拉取 Dify 官方镜像并完成部署：配置 docker-compose.yml（包含 Dify API Server/Worker/Web/PostgreSQL/Redis 共 5 个服务），"
            "启动后在 Web 界面完成了大模型 API Key 配置（接入 DeepSeek 和通义千问两个模型供应商）。\n\n"
            "搭建 Dify 知识库：上传项目文档（PDF/Word/TXT 多种格式），设置分段规则（自动检测段落边界 + 最大 1000 tokens/段），"
            "选择 Embedding 模型并构建索引。进行了多轮真实召回测试——输入与文档相关的问题，验证能否正确召回对应的原文段落，并对分段策略做了调优。\n\n"
            "编排智能客服工作流：问题分类节点（LLM 判断意图）→ 知识库检索节点（条件路由：有命中→走 LLM 总结，无命中→走安全转交）"
            "→ LLM 回复节点 → TTS 文字转语音节点（支持语音播报）→ 文生图节点（可选）。工作流发布后通过 Dify API 在项目中实际调用验证。\n\n"
            "关键架构决策：FastAPI 不直接持有 Dify 密钥，而是通过统一网关层转发请求，前端完全不接触第三方密钥。当 Dify 服务不可用时自动降级到本地 RAG/Agent。"
        ),
        "think": (
            "Dify 这类低代码平台的价值在于「降低 AI 应用落地门槛」——拖拽式编排、可视化调试、开箱即用的日志与版本管理，"
            "让非算法背景的开发者也能快速搭建可用的 AI 工作流。但它与 LangGraph 并非替代关系：Dify 适合标准化流程（如 FAQ 机器人），"
            "而 LangGraph 适合需要精细控制的复杂编排（如多步推理、动态路由）。两者互补使用效果最好。\n\n"
            "统一网关层的设计让我意识到「安全解耦」的重要性——密钥集中管理、服务隔离、失败自动降级，这些在生产环境中缺一不可。"
        ),
    },
    {
        "d": "7月23日", "date": "2026年7月23日",
        "work": (
            "全栈集成联调日——将前面开发的各个模块串成完整产品。\n\n"
            "前端侧：完善了三个核心交互界面——（1）AI 对话面板：消息列表滚动、Markdown 渲染、代码高亮、复制按钮；（2）"
            "知识库问答页：搜索框、结果列表卡片（展示摘要与相关度分数）、点击展开全文及引用来源；（3）"
            "客服聊天窗口：实时 SSE 消息推送、语音播放按钮（调用 TTS API）、文件上传预览。\n\n"
            "后端侧：FastAPI 作为统一入口，根据请求类型路由分发——普通对话走原生流式 LLM、知识问答走 RAG Pipeline、"
            "复杂任务走 LangGraph Agent 工作流、标准化流程走 Dify 网关。引入 Redis 缓存层：对高频重复查询（相同用户+相同问题+相同偏好设置）"
            "缓存 LLM 返回结果（TTL 10 分钟），减少约 40% 的重复 API 调用。持久化了对话历史记录（MySQL）、知识库文档管理、用户偏好设置。\n\n"
            "全天反复进行端到端联调——从前端发请求到后端路由到 AI 响应渲染的全链路，修复了字段命名不一致、时间格式未统一、SSE 断连重试等多个集成问题。"
        ),
        "think": (
            "全栈联调是最容易出问题的阶段——每个模块单独跑得好好的，串起来就各种不兼容。今天的经验总结：\n"
            "(1) 接口契约必须前置约定（Request/Response 的 JSON Schema），不能靠口头沟通；\n"
            "(2) 前后端并行开发时 Mock 服务必不可少；\n"
            "(3) 缓存键的设计要覆盖用户 ID + 上下文 + 偏好维度，否则不同用户会拿到串味的结果；\n"
            "(4) 「客户端取消即回滚、异常即降级」的容错策略要提前明确，否则界面状态和数据很容易不一致。"
        ),
    },
    {
        "d": "7月24日", "date": "2026年7月24日",
        "work": (
            "容器化部署与 DevOps 实践。\n\n"
            "编写 Dockerfile：后端基于 `python:3.12-slim` 镜像，安装依赖后用 uvicorn 启动；前端基于 `node:22-alpine`，"
            "npm run build 后用 nginx 托管静态文件；Redis 直接使用官方镜像。编写 docker-compose.yml 将三个服务编排在一起，"
            "`docker compose up -d` 一条命令即可在任意 Linux 机器上拉起整套环境。\n\n"
            "入门 Kubernetes 基础概念：Pod（最小调度单元）、Service（集群内部负载均衡与服务发现）、Ingress（七层路由，"
            "配置域名与路径分发规则）。了解了 Deployment 的滚动更新策略（RollingUpdate）与健康检查探针（Liveness/Readiness Probe）。\n\n"
            "配置 CI/CD 流水线：代码提交触发 GitHub Actions → 自动运行 pytest 单元测试 + Postman 接口集合测试 → 构建 Docker 镜像并推送到镜像仓库"
            "→ SSH 连接目标服务器拉取最新镜像并重启容器。在线上环境完成了完整的功能回归测试和 AI 智能体端到端验证。"
        ),
        "think": (
            "容器化解决了经典口号「在我机器上能跑」的问题——环境被固化进镜像，任何机器都能一致运行。"
            "Docker Compose 让多服务编排变得声明式且可版本控制，Kubernetes 则把这个能力扩展到了生产级规模（虽然我们目前只用到了最基础的 Pod/Service/Ingress）。\n\n"
            "CI/CD 自动化的价值不仅是省人力，更重要的是「可复现性」——每次部署走的都是同一条流水线，排除了人为操作的随机错误。"
            "云上运维和本地开发差异确实很大（网络延迟、权限模型、日志收集），提前做好健康检查和结构化日志能在故障发生时节省大量排查时间。"
        ),
    },
    {
        "d": "7月25日", "date": "2026年7月25日",
        "work": (
            "项目收尾与答辩准备。\n\n"
            "整理了一份完整的功能演示脚本，确保演示过程流畅有序：\n"
            "① 登录企业用户账号 → 进入仪表盘看数据概览；\n"
            "② 打开 AI 对话窗口提问（观察 SSE 流式输出效果）；\n"
            "③ 切换到知识库问答，输入与文档相关的问题（查看引用来源与融合评分）；\n"
            "④ 新建客服工单（观察实时分类建议与自动填充）；\n"
            "⑤ 切换管理员账号，调整检索参数权重（即时生效）；\n"
            "⑥ 切换决策者角色，浏览运营看板（对话量/满意度/热门话题统计）。\n\n"
            "准备答辩 PPT：按照「项目背景 → 需求分析 → 技术选型 → 构构设计 → 核心功能展示（附代码讲解）"
            "→ 工程亮点 → 部署方案 → 总结展望」的结构组织内容，重点准备了三处核心技术点的讲解："
            "(1) 原生 SSE 流式转发实现；(2) 混合 RAG 加权 RRF 融合检索；(3) LangGraph 多智能体 StateGraph 编排。\n\n"
            "对照实训考核八项（功能完整性、代码规范性、文档完备性、创新点、团队协作、答辩表现、工程实践、安全意识）逐项自查确认。"
        ),
        "think": (
            "九天实训下来，真正打通了从「调一个 API」到「交付一套完整 AI 应用系统」的全链路。回顾整个历程：\n\n"
            "技术层面——掌握了 FastAPI/Vue3 全栈开发、大模型 API 对接与流式处理、LangChain/LangGraph 框架应用、"
            "RAG 检索增强、Dify 低代码平台集成、Docker/K8s 容器化部署这一套完整的 AI 工程栈。\n\n"
            "工程层面——体会到了规范文档、接口契约、自动化测试、CI/CD 这些「非功能性工作」对项目长期可维护性的决定性作用。\n\n"
            "思维层面——AI 辅助开发已经从「尝鲜」变成「日常」，但核心原则不变：提示要精准、输出要审核、优化要有据。"
            "这次实训不仅是技能积累，更是对「AI 时代软件工程师应该具备什么能力」的一次系统性认知升级。"
        ),
    },
]

OUT = r"D:\shixi\实习日记"
os.makedirs(OUT, exist_ok=True)

for i, e in enumerate(entries, 1):
    doc = Document()

    # ---- 标题 ----
    para(doc, "暨南大学本科生实习日记", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    para(doc, "（集中实习用）", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)

    # ---- 信息表格（严格对齐截图模板）----
    make_header_table(doc, e["date"])

    para(doc, "", after=6)  # 表格后留白

    # ---- 正文 ----
    para(doc, e["date"], size=12, bold=True, after=4)
    para(doc, "一、当天实习工作主要内容：", size=12, bold=True, after=2)
    para(doc, e["work"], size=12, indent=True, after=10)
    para(doc, "二、当天实习工作的思考与收获：", size=12, bold=True, after=2)
    para(doc, e["think"], size=12, indent=True, after=10)
    para(doc, "实习生(手写签名)：________________", size=12, before=12, after=0)

    fn = f"第{i}天_{e['d']}.docx"
    doc.save(os.path.join(OUT, fn))
    print("saved", fn)

print("ALL DONE, total", len(entries))
