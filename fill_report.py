from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import _Cell


TEMPLATE = Path(r"C:\Users\大帝之资\xwechat_files\wxid_vn11ora3b0oe12_b139\msg\file\2026-07\XXX项目报告模板(1).docx")
OUTPUT = Path(r"D:\shixi\XXX项目报告_邓子川.docx")


def set_paragraph(paragraph, text):
    """Replace visible text while retaining the paragraph and first-run format."""
    runs = list(paragraph.runs)
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run._element.getparent().remove(run._element)
    elif text:
        paragraph.add_run(text)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def set_cell(cell, text):
    if not cell.paragraphs:
        cell.add_paragraph()
    set_paragraph(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        set_paragraph(paragraph, "")


def set_tc(table, tc, text):
    """Set a physical XML cell, avoiding python-docx merged-cell aliases."""
    set_cell(_Cell(tc, table), text)


def main():
    doc = Document(str(TEMPLATE))
    paragraphs = doc.paragraphs

    replacements = {
        15: "东软智慧商务AI助手平台的设计与实现",
        16: "摘要",
        17: "传统商务服务依赖人工查找制度和重复答复，客户等待时间长，客服难以保持统一口径，管理者也难以及时发现高频问题。本项目设计并实现东软智慧商务AI助手平台，采用Vue 3与FastAPI构建前后端分离系统，以本地混合RAG和受控LangGraph Agent提供可离线验证的智能问答，并以OpenAI-compatible模型和Dify作为可选增强。平台面向企业用户、客服人员、系统管理员和决策者，提供智能对话、知识库检索、客服工单、实时事件、AI配置与运营看板等功能。系统通过引用、Agent轨迹、人工确认和异常降级保证回答可解释、可追溯和可控，现有自动化验收和固定评测结果表明核心链路可以稳定运行。",
        18: "本报告按照背景、问题、方案、实现、测试和总结的逻辑展开，重点记录平台的需求分析、总体架构、AI编排、全栈实现及部署验证。",
        19: "",
        20: "",
        21: "",
        22: "",
        23: "",
        24: "",
        25: "",
        26: "",
        27: "",
        28: "",
        29: "关键词：智慧商务；AI助手；FastAPI；Vue 3；混合RAG；LangGraph",
        31: "ABSTRACT",
        32: "Traditional business services rely heavily on forms, manual policy lookup and repeated replies, which leads to long customer waiting time and inconsistent support quality. This project designs and implements the Neusoft Smart Business AI Assistant Platform. Vue 3 and FastAPI are used to build a separated front-end and back-end system, while an offline hybrid RAG pipeline and a controlled LangGraph agent provide reproducible AI capabilities. OpenAI-compatible models and Dify are integrated as optional enhancement services. The platform serves enterprise users, support agents, administrators and executives with intelligent chat, knowledge retrieval, support tickets, real-time events, AI configuration and an operation dashboard. Citations, agent traces, human confirmation and explicit fallback paths make the responses explainable, traceable and controllable. The existing automated acceptance records and fixed evaluation set show that the core workflow runs reliably.",
        33: "KEY WORDS: Smart business AI assistant; FastAPI; Vue 3; Hybrid RAG; LangGraph",
        63: "",
        65: "商务服务场景中存在大量重复咨询、制度检索和工单流转工作。传统方式主要依赖人工经验，既影响响应速度，也容易产生口径不一致和信息遗漏。随着大模型、检索增强生成和智能体编排技术逐渐成熟，将企业知识、客服流程与统一对话入口结合，能够为业务人员提供更高效的辅助工具。",
        66: "本项目聚焦企业用户、客服人员、系统管理员和决策者四类角色，围绕“可运行、可解释、可回退、可人工接管”建设一套教学演示级智慧商务AI平台。",
        67: "平台不把模型生成结果直接视为事实，而是通过知识引用、处理轨迹、工具白名单和人工确认约束风险；没有外部模型密钥时，核心问答仍可由本地链路完成。",
        68: "项目目标是打通从前端交互、后端业务、知识检索、智能体编排到容器部署的完整AI应用工程链路。",
        69: "报告内容与代码仓库中的需求、架构、AI设计、测试和部署文档保持一致，所有已执行结果与尚未具备的生产证据分别说明。",
        71: "企业知识通常分散在制度、FAQ、合同与操作手册中，客服需要在多个文档之间反复查找，企业用户也难以获得及时、统一的答复。大模型可以改善自然语言交互，但直接生成容易出现知识过期、事实幻觉和敏感业务承诺等问题。因此，本项目采用检索增强生成，将文档切分后建立本地向量索引和关键词索引，并在回答中返回引用片段，使用户可以核验答案来源。",
        73: "国内企业智能客服研究主要集中在知识库问答、工单自动分类和大模型辅助客服等方向。常见方案已经能够完成意图识别、FAQ检索和推荐回复，但在离线复现、引用可追溯、工具安全边界以及异常回退方面仍有不足。本项目将知识检索、客服工单和运营分析放在同一业务平台中，强调统一接口、角色权限和人工在环。",
        75: "国外相关产品和研究更重视对话式搜索、企业知识治理、Agent工具调用和可观测性。先进方案通常提供引用、日志、权限控制和多步骤流程编排，但往往依赖云端模型、托管向量数据库或复杂的企业集成环境。项目在吸收这些思路的同时，保留本地确定性Embedding、FAISS和SQLite回退链路，适合课堂环境验证和后续替换为生产组件。",
        77: "本课题旨在设计并实现一套面向企业商务服务的AI助手平台，降低重复咨询和知识查找成本，提高客服响应与管理分析效率，同时避免模型在高风险问题上臆造业务事实。",
        78: "研究重点包括：构建前后端分离的业务平台；实现带引用和轨迹的智能问答；建立可解释、可回退的混合RAG与LangGraph编排链路；完成工单、管理、看板、容器部署和测试验证。",
        79: "平台采用Vue 3 + Vite前端、FastAPI后端、SQLAlchemy数据访问和SQLite演示数据库，Redis用于缓存，Docker Compose与Kubernetes清单用于部署。",
        80: "具体研究内容如下：",
        81: "（1）分析企业用户、客服、管理员和决策者的使用场景，明确功能需求、权限边界和非功能需求；",
        82: "（2）设计前后端分离架构、业务数据模型、知识库索引和AI请求数据流；",
        83: "（3）实现原生SSE流式输出、混合RAG、受控多智能体、客服工单、偏好与缓存等核心功能；",
        84: "（4）通过自动化测试、固定AI评测、Compose健康检查和浏览器冒烟验证系统的可运行性与稳定性。",
        88: "系统分析阶段围绕“做什么”和“能不能做”展开。项目以四类角色为边界，以问答、知识、工单、管理和看板为核心业务闭环，并将模型异常、权限越权、敏感问题和部署差异纳入分析范围。",
        90: "平台的可行性从技术、经济、操作和法律/社会四个维度进行论证。",
        91: "技术可行性：Vue 3、FastAPI、SQLAlchemy、LangChain、LangGraph、FAISS、Redis、Docker等技术成熟且文档完善。项目通过本地确定性Embedding和SQLite回退降低外部依赖，OpenAI-compatible与Dify只位于适配边界，便于替换和测试。",
        92: "经济可行性：项目以教学演示为目标，采用开源框架和本地运行方式，无需购买专用硬件或固定云服务。Docker Compose可以一次启动前端、后端和Redis，开发与演示成本主要是实训时间和普通开发电脑资源。",
        93: "操作可行性：四类角色均有明确入口和权限。企业用户通过自然语言提问，客服在工作台中查看工单并确认回复，管理员维护知识和AI设置，决策者查看运营看板。前端提供引用、状态和回退标识，降低用户理解成本。",
        94: "法律与社会可行性：密钥只通过环境变量管理，受保护接口使用Bearer Token和角色校验；工具只允许读取工单聚合，不暴露客户明细，也不执行任意写操作。合同、付款、订单和故障等高风险问题要求补充信息或转人工。生产使用时仍需补充隐私、脱敏、审计和数据授权制度。",
        96: "系统面向企业商务服务，支持企业用户获取业务帮助、客服处理和确认工单、管理员维护知识与策略、决策者分析服务数据。前端通过相对路径访问统一API，后端根据角色和请求类型选择业务服务或AI编排链路。",
        97: "企业用户可登录、发起多轮问答、查看历史会话、检索知识并提交咨询工单；客服人员可查看待处理工单、使用AI推荐回复、保存草稿并显式确认结案。",
        98: "管理员可管理用户、知识文档、检索参数、系统提示词和审计消息；决策者可查看咨询量、工单状态、问题分类、质检代理指标和分析报告。",
        99: "AI问答流程可抽象为：接收问题→意图分类与知识检索并行→白名单工具路由→生成回答→引用和轨迹整理→质量检查与必要回退。",
        100: "非功能需求包括：无外部密钥时可运行；回答带引用和处理轨迹；受保护接口返回401/403；流式取消后不残留未完成消息；缓存按用户和上下文隔离；演示规模下检索与回答在数秒内完成。",
        102: "功能性需求按角色归纳如下：企业用户负责问答、历史会话、知识检索、偏好设置和工单提交；客服人员负责工单队列、AI建议、草稿及人工确认；管理员负责用户、知识、提示词、检索设置和审计；决策者负责运营指标和分析报告。",
        103: "每项功能均具有可验证输入、处理逻辑和输出。例如知识库上传需校验格式、大小和有效文本长度；保存回复草稿不得自动改变工单状态；只有显式提交resolved才代表人工确认完成。",
        105: "性能需求：演示数据规模下本地检索和回答应在数秒内返回，知识库单文件上传上限为5 MB，检索Top-K可配置。安全需求：Token认证、角色授权、密钥不入库、工具白名单和参数校验。",
        106: "易用性与可靠性需求：界面提供加载、错误、回退和引用状态；浏览器不支持语音时退回文本流程；模型首Token前失败使用本地回退，部分Token后失败通过reset保证界面、完成事件和数据库内容一致。",
        109: "总体设计把平台划分为前端层、API与业务层、AI编排层、数据与缓存层以及部署运维层，各层通过稳定边界协作。核心原则是离线优先、接口稳定、证据可见、人工兜底和渐进演进。",
        110: "前端使用Vue 3、Vue Router、Pinia、Axios和Vite；后端使用FastAPI、Pydantic和SQLAlchemy；SQLite保存演示业务数据，Redis保存检索与最终回答缓存；AI层组合LangChain/LangGraph、本地Embedding、FAISS、OpenAI-compatible适配器和Dify网关。",
        111: "模块按角色和能力划分为智能助手、知识库、客服工作台、管理后台、运营看板、用户偏好、实时工单事件和媒体交互，各模块均通过/api/v1统一访问。",
        112: "数据库实体包括users、user_preferences、conversations、messages、knowledge_documents、support_tickets和ai_settings；AI回答中的引用、轨迹、回退标记与消息记录关联保存。",
        114: "系统采用B/S前后端分离架构。Vue 3负责页面、路由、状态和SSE消费；FastAPI负责认证、业务CRUD、权限、流式适配和AI编排；AI编排层执行分类、检索、工具和质量检查；SQLite/Redis提供持久化与缓存。配置OpenAI-compatible模型或Dify后，统一API可按环境变量启用远程增强；未配置或失败时回退到本地链路。",
        116: "平台功能模块由“用户入口、知识与客服、管理与分析、AI基础能力”四组组成。",
        117: "用户入口包含登录、智能助手、会话历史、偏好设置、语音输入/朗读和工单申请；知识与客服包含文档维护、检索测试、工单队列、AI建议和人工确认；管理与分析包含用户管理、AI设置、消息审计、指标看板和报告；AI基础能力包含RAG、LangGraph、缓存、Dify网关和降级机制。",
        118: "角色权限与模块一一对应：enterprise_user不能访问管理员和决策者页面，support_agent只处理客服范围，admin维护系统配置，executive只读运营分析。服务端再次校验权限，前端路由控制仅用于改善体验。",
        119: "模块设计遵循单一职责，前端页面通过API客户端调用后端，后端服务层封装数据与AI逻辑，便于测试和替换外部模型。",
        121: "数据库采用关系模型，核心实体围绕用户、会话、消息、知识文档、工单和系统设置组织。文档与知识片段、会话与消息、用户与工单均保持可追溯关系。",
        122: "概念设计实体包括用户、用户偏好、会话、消息、知识文档、工单和AI设置；用户与会话、工单是一对多关系，会话与消息是一对多关系，知识文档与知识片段是一对多关系。",
        123: "逻辑设计使用SQLite表保存演示数据，字段由Pydantic模型和SQLAlchemy模型共同约束。AI引用与轨迹作为消息的结构化字段保存，避免回答和证据脱离。",
        125: "用户实体保存邮箱、密码哈希、角色和启停状态；偏好实体保存回答风格、语言和自动朗读开关；会话和消息实体保存多轮交互；知识文档保存标题、来源、状态和索引版本；工单保存问题、分类、优先级、建议回复和最终回复；AI设置保存提示词、Top-K和分块参数。",
        126: "主要实体及属性：users(id、email、role、is_active)；user_preferences(user_id、response_style、preferred_language、auto_play_voice)；conversations(id、user_id、title、status)；messages(id、conversation_id、role、content、citations、trace)；knowledge_documents(id、title、source、status、version)；support_tickets(id、requester_id、category、priority、status、suggested_reply、final_reply)。",
        127: "关系约束：一个用户可以拥有多个会话和工单；一个会话包含多条消息；一个知识文档可以产生多个索引片段；工单的requester_id关联用户；管理员设置为全局配置，由受保护接口维护。",
        128: "E-R关系图应体现用户—会话—消息、用户—工单、知识文档—知识片段以及管理员—AI设置等关系，后续逻辑表和API字段与该关系保持一致。",
        129: "为保证数据安全，系统只向AI工具提供工单队列的聚合数量，不返回客户名称、问题正文、工单编号等明细；高风险业务不通过模型直接写入数据库。",
        130: "数据库概念设计以“业务实体清晰、关系可追溯、权限边界明确”为原则，为后续接口和测试提供统一依据。",
        133: "逻辑结构设计以实体关系为基础，将用户、偏好、会话、消息、知识、工单和设置映射为关系表。主键采用整数或字符串标识，外键用于维护用户、会话与工单的归属，状态字段使用受控枚举。",
        134: "关键表结构如下：users保存认证与角色；user_preferences保存个性化设置；conversations/messages保存对话；knowledge_documents/knowledge_chunks保存知识原文与索引；support_tickets保存工单状态流转；ai_settings保存管理员配置。",
        135: "字段设计重点是可验证性和最小权限：密码只保存哈希，消息引用与轨迹采用结构化JSON，知识文档状态控制是否参与检索，工单status区分open、in_progress和resolved，final_reply不自动触发结案。",
        136: "表之间通过user_id、conversation_id、requester_id和document_id关联。删除知识文档时同时清理对应片段；会话和工单查询按当前用户或角色过滤，避免跨用户数据泄露。",
        137: "数据库结构与后端模型、API响应和前端类型保持一致，测试覆盖正常、401、403、422、状态流转和缓存隔离等边界。",
        142: "详细设计围绕智能问答、知识检索、客服工单和运营分析等核心模块展开，说明其输入、处理、输出及异常策略。",
        143: "每个模块均有明确服务边界：API层负责认证和参数校验，业务层负责数据操作，AI层负责分类、检索、工具和回答计划，前端负责状态呈现与用户交互。",
        144: "设计重点包括原生SSE时序、RAG混合排序、LangGraph条件路由、工单SSE事件和草稿/结案状态约束。",
        145: "详细设计与总体架构、需求角色和数据库关系保持一致，避免把未实现的生产能力写成已完成事实。",
        147: "智能问答流程：开始→校验Token和请求→保存用户消息→分类与知识检索并行→条件路由白名单工具→组织带证据的回复计划→质量检查→模型原生流或本地回答→发送done并保存助手消息。",
        148: "知识库流程：上传文件→检查扩展名、大小和有效文本→段落切分→确定性Embedding→写入向量与关键词索引→返回ready状态和分块数量。",
        149: "客服工单流程：企业用户提交问题→创建open工单→后台生成分类和建议回复→客服保存草稿→人工修改→显式提交resolved→向受保护SSE订阅者推送状态更新。",
        150: "异常分支包括模型首Token前回退、部分Token后reset、客户端取消回滚、无证据时转人工以及Dify不可用时切换本地RAG/Agent。",
        151: "流程设计遵循输入、判断、处理、输出四类节点，所有外部模型调用均通过统一网关，便于替换、测试和审计。",
        153: "核心对象包括AuthService、KnowledgeService、BusinessAgentOrchestrator、TicketService、DifyGateway和CacheService。",
        154: "AuthService负责Token签发与角色校验；KnowledgeService负责文档解析、切分、索引和检索；BusinessAgentOrchestrator负责StateGraph状态、分类、检索、工具和质检；TicketService负责工单状态和实时事件；DifyGateway负责远程工作流调用与降级；CacheService负责键生成、TTL和用户隔离。",
        155: "对象之间通过接口调用而非直接共享页面状态。AI编排器只接收经过校验的请求和当前会话上下文，工具执行器只暴露白名单聚合能力，避免模型获得任意数据库写权限。",
        156: "前端对应对象包括auth store、assistant API、knowledge API、ticket API和dashboard API；后端Pydantic schema保证请求与响应字段稳定。",
        157: "核心对象的职责、数据库实体和功能模块一一对应，便于进行单元测试、集成测试和后续生产替换。",
        159: "项目的关键算法主要是混合RAG排序和流式协议处理。文档先按约500字符、50字符重叠切分，向量检索与BM25关键词检索并行执行，再按向量0.65、关键词0.35进行加权RRF融合；排序结果返回Top-K引用。原生模型流按SSE事件解析delta、[DONE]和tool_calls，首Token立即转发，客户端取消关闭上游流。",
        161: "系统实现",
        162: "企业用户、客服与管理角色模块",
        163: "系统实现以需求分析中的角色和模块为依据，重点展示智能对话、知识库检索、客服工单、管理配置和运营看板。代码采用Vue 3 + FastAPI分层结构，外部模型与Dify均通过网关接入。",
        164: "智能对话与知识库问答",
        165: "企业用户登录后可在智能助手中进行多轮问答，也可以切换知识库模式。后端先保存用户消息，再由LangGraph编排分类、混合检索和受控工具，最终返回回答、引用、trace和used_fallback标记。配置模型时使用stream=true消费真实SSE delta；未配置或失败时回退本地确定性链路。知识库支持文本、Markdown、CSV、PDF和DOCX解析，返回文档标题、片段与融合相关度。",
        167: "客服工单与实时会话",
        168: "企业用户可提交问题和优先级，系统创建工单并异步生成分类与建议回复。客服工作台通过受保护SSE接收创建和更新事件，先保存回复草稿，再显式确认结案。若AI建议需要人工处理，企业用户可发起转人工会话，客服和用户通过各自权限范围内的实时消息接口沟通。",
        169: "该模块强调人工在环：模型不能自动发送承诺性回复，工单草稿不会隐式改变状态，客服确认后才写入resolved。服务端按用户归属过滤事件，防止跨用户泄露。",
        170: "管理配置与运营看板",
        171: "系统测试",
        172: "测试覆盖认证授权、智能问答、知识库、工单、管理、看板、SSE流式时序、缓存隔离、Dify网关和容器部署。测试结果只引用项目中实际执行并有退出码的记录，未将静态配置或计划误写成线上部署。",
        174: "测试目的",
        175: "验证系统是否满足功能需求和非功能需求，重点检查角色权限、核心业务闭环、回答引用与轨迹、流式取消和回退、工单状态约束、知识索引可用性以及部署健康检查。",
        176: "测试方法",
        177: "采用白盒单元测试、FastAPI接口集成测试、固定AI评测、容器健康检查和浏览器冒烟相结合的方法。固定离线评测共12条题目，检查分类、Top-3检索、引用、Agent轨迹和安全转人工；前端和后端核心路径按角色执行。",
        178: "功能测试",
        179: "登录、智能问答、知识检索与工单闭环测试",
        180: "功能测试覆盖企业用户登录和问答、管理员知识库维护、客服工单草稿与显式结案、决策者看板以及越权访问。验收记录显示后端全量测试85 passed，固定AI评测12/12达到阈值，Dify/媒体定向测试21 passed，课程Compose服务和前端代理健康检查通过。测试同时验证模型流式首Token时序、客户端取消后不提交未完成消息、缓存按用户隔离以及无证据时转人工。",
        181: "表6.1 核心功能测试用例",
        183: "非功能性测试",
        184: "性能与稳定性：演示数据规模下本地检索和回答在数秒内完成；模型未配置、超时或异常时返回明确回退状态；Redis不可用时使用有界内存缓存。",
        185: "安全性：受保护API使用Bearer Token，管理员接口校验角色；上传限制大小和格式；工具仅允许白名单参数和只读工单聚合，不执行任意SQL或业务写操作；密钥不进入仓库。",
        186: "可维护性与兼容性：接口提供OpenAPI文档，前端以相对路径访问API；Docker Compose和Kubernetes清单可解析，镜像以非root用户运行并配置健康探针。生产集群rollout、GHCR推送和真实业务资料评测不在本次报告的已完成范围内。",
        187: "总结与展望",
        188: "本次实训完成了从需求分析、架构设计、全栈实现、AI能力接入到测试和容器运行的完整链路。项目掌握了Vue 3/FastAPI开发、原生SSE、混合RAG、LangGraph受控编排、Dify网关、Redis缓存、Docker/Kubernetes和CI配置等技术，并通过引用、轨迹、回退和人工确认建立了可解释的业务边界。",
        189: "项目仍属于教学演示系统，后续可将SQLite替换为PostgreSQL，将本地Embedding/FAISS替换为托管向量数据库，引入Redis Pub/Sub支持多实例工单事件，并补充真实业务资料、隐私脱敏、监控告警和生产网络实测。对于模型质量，应持续使用人工复核和真实业务评测，而不能把固定题集结果等同于客户满意度。",
        191: "[1] 项目组：《需求规格说明书》，D:\\shixi\\docs\\requirements.md，2026。\n[2] 项目组：《总体架构设计》，D:\\shixi\\docs\\architecture.md，2026。\n[3] 项目组：《AI设计说明》，D:\\shixi\\docs\\ai-design.md，2026。\n[4] 项目组：《测试与AI评估计划》，D:\\shixi\\docs\\test-evaluation.md，2026。\n[5] 项目组：《项目验收执行报告》，D:\\shixi\\docs\\acceptance-report.md，2026。",
    }

    for index, text in replacements.items():
        set_paragraph(paragraphs[index], text)

    # Remove example-project drawings while preserving the surrounding paragraph geometry.
    for index in (131, 138, 166, 168, 169):
        clear_paragraph(paragraphs[index])

    # Cover metadata.
    cover = doc.tables[0]
    # The cover uses merged cells; address the physical cells in each row.
    cover_cells = [row._tr.tc_lst for row in cover.rows]
    set_tc(cover, cover_cells[0][0], "信息科学技术")
    set_tc(cover, cover_cells[1][0], "计算机科学与技术")
    set_tc(cover, cover_cells[2][1], "邓子川")
    set_tc(cover, cover_cells[3][1], "2023104525")
    set_tc(cover, cover_cells[4][1], "2026年7月17日—25日")

    # Keep the scoring table, but make it match the actual project and leave the grade column for the instructor.
    score = doc.tables[1]
    set_cell(score.cell(1, 1), "完成项目需求、架构和相关资料分析")
    set_cell(score.cell(2, 1), "完成四类角色需求分析并绘制功能结构说明")
    set_cell(score.cell(3, 1), "完成用户、会话、消息、知识、工单等实体关系设计")
    set_cell(score.cell(4, 1), "完成SQLite逻辑表、索引、状态和外键关系设计")
    set_cell(score.cell(5, 1), "完成Vue 3前端页面与交互界面")
    set_cell(score.cell(6, 1), "完成问答、RAG、Agent、工单、管理和看板功能")
    set_cell(score.cell(7, 1), "完成实训报告撰写")
    set_cell(score.cell(8, 1), "报告结构和模板格式保持一致")
    set_cell(score.cell(9, 1), "需求、设计、实现和测试逻辑一致")

    # Test-case table from the template, with the same geometry and 6-column layout.
    test = doc.tables[2]
    rows = [
        ["项目名称", "东软智慧商务AI助手平台", "东软智慧商务AI助手平台", "程序版本", "V1.0", "V1.0"],
        ["测试环境", "Windows 11；Python 3.12；Vue 3；FastAPI；SQLite；Docker Compose", "Windows 11；Python 3.12；Vue 3；FastAPI；SQLite；Docker Compose", "Windows 11；Python 3.12；Vue 3；FastAPI；SQLite；Docker Compose", "Windows 11；Python 3.12；Vue 3；FastAPI；SQLite；Docker Compose", "Windows 11；Python 3.12；Vue 3；FastAPI；SQLite；Docker Compose"],
        ["功能模块", "智能对话与知识库问答", "智能对话与知识库问答", "智能对话与知识库问答", "智能对话与知识库问答", "智能对话与知识库问答"],
        ["编制人", "邓子川", "邓子川", "编制时间", "2026-07-25", "2026-07-25"],
        ["功能特性", "验证登录、问答、引用、Agent轨迹和本地回退是否正确", "验证登录、问答、引用、Agent轨迹和本地回退是否正确", "验证登录、问答、引用、Agent轨迹和本地回退是否正确", "验证登录、问答、引用、Agent轨迹和本地回退是否正确", "验证登录、问答、引用、Agent轨迹和本地回退是否正确"],
        ["测试目的", "验证核心问答链路在正常、无证据和模型不可用场景下可解释、可回退", "验证核心问答链路在正常、无证据和模型不可用场景下可解释、可回退", "验证核心问答链路在正常、无证据和模型不可用场景下可解释、可回退", "验证核心问答链路在正常、无证据和模型不可用场景下可解释、可回退", "验证核心问答链路在正常、无证据和模型不可用场景下可解释、可回退"],
        ["预置条件", "服务健康，演示数据库和知识文档已初始化，企业用户可登录", "服务健康，演示数据库和知识文档已初始化，企业用户可登录", "服务健康，演示数据库和知识文档已初始化，企业用户可登录", "服务健康，演示数据库和知识文档已初始化，企业用户可登录", "服务健康，演示数据库和知识文档已初始化，企业用户可登录"],
        ["参考信息", "requirements.md；ai-design.md", "无", "特殊说明", "无", "无"],
        ["用例编号", "输入数据", "操作步骤", "预期输出", "测试结果", "用例说明"],
        ["AI-001", "账号：enterprise@neusoft.local\n问题：发票申请需要准备什么材料？", "登录后进入智能助手并发送问题", "返回回答、引用、分类/检索/回复轨迹", "通过", "验证有知识证据时的问答链路"],
        ["AI-002", "问题：今天天气如何？", "在智能助手中发送无业务依据的问题", "说明知识库范围，不编造业务结论并建议人工核验", "通过", "验证无证据安全回退"],
        ["AI-003", "问题：服务中断影响客户演示，怎么办？", "发送故障类问题并查看结果", "给出升级建议，要求补充影响信息，不承诺修复时间", "通过", "验证高风险问题人工兜底"],
        ["AI-004", "管理员上传TXT并输入合同审批需要多久？", "上传文档、重建索引、执行检索测试", "文档进入ready状态并返回相关片段和分数", "通过", "验证知识维护和检索"],
        ["AI-005", "客户端在首Token后取消请求", "调用流式接口并关闭客户端连接", "关闭上游流，不提交未完成助手消息；必要时发送reset", "通过", "验证原生SSE取消与一致性"],
        ["AI-006", "企业用户提交工单后客服保存回复草稿", "创建工单、编辑回复并保存，再显式提交结案", "草稿保持in_progress，显式确认后才变为resolved", "通过", "验证人工在环和状态约束"],
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            set_cell(test.cell(r, c), value)

    # Save a new document; the retained template remains untouched.
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
