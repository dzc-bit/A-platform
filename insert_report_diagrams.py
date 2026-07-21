from pathlib import Path
from docx import Document
from docx.shared import Inches


SOURCE = Path(r"D:\shixi\XXX项目报告_邓子川.docx")
OUTPUT = Path(r"D:\shixi\XXX项目报告_邓子川_含图.docx")
IMG = Path(r"D:\shixi\report_diagrams")


def add_picture_before(paragraph, path):
    new_paragraph = paragraph.insert_paragraph_before()
    new_paragraph.alignment = paragraph.alignment
    run = new_paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.1))
    return new_paragraph


def main():
    doc = Document(str(SOURCE))
    p = doc.paragraphs
    add_picture_before(p[115], IMG / "architecture.png")
    add_picture_before(p[120], IMG / "modules.png")
    add_picture_before(p[132], IMG / "er.png")
    add_picture_before(p[152], IMG / "flow.png")
    add_picture_before(p[158], IMG / "classes.png")
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
